#!/usr/bin/env python3
"""Generate Big-4 consulting-style DOCX from .pbi/doc_data.json.

Usage: python gen_docx.py <doc_data.json> <output.docx>
Requires: pip install python-docx
"""
import sys
import json
from datetime import datetime

NAVY  = (0x1B, 0x3A, 0x6B)
GOLD  = (0xC9, 0xA8, 0x4C)
GREY  = (0xF2, 0xF2, 0xF2)
DGREY = (0x40, 0x40, 0x40)
WHITE = (0xFF, 0xFF, 0xFF)
LGREY = (0xD9, 0xD9, 0xD9)
CGREEN = (0x1E, 0x6A, 0x3A)


def check_import():
    try:
        from docx import Document  # noqa
        return True
    except ImportError:
        print('MISSING_DEP python-docx  —  run: pip install python-docx', file=sys.stderr)
        sys.exit(2)


def rgb(t):
    from docx.shared import RGBColor
    return RGBColor(*t)


def set_para(p, bold=False, size=10, color=DGREY, align='left', sb=0, sa=6):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    al = {'left': WD_ALIGN_PARAGRAPH.LEFT,
          'center': WD_ALIGN_PARAGRAPH.CENTER,
          'right': WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = al.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    for run in p.runs:
        run.bold = bold
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.color.rgb = rgb(color)
    return p


def add_para(doc, text, bold=False, size=10, color=DGREY, align='left', sb=0, sa=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Calibri'
    return set_para(p, bold=bold, size=size, color=color, align=align, sb=sb, sa=sa)


def add_heading(doc, number, title, level=1):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    sizes = {1: 17, 2: 13, 3: 11}
    sbs   = {1: 20, 2: 14, 3: 8}
    size = sizes.get(level, 11)
    sb   = sbs.get(level, 8)
    text = f'{number}  {title}' if number else title
    p = add_para(doc, text, bold=True, size=size, color=NAVY, sb=sb, sa=4)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'C9A84C')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove existing shd
    for existing in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tc_pr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    from docx.shared import Pt, Cm, RGBColor
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = 'Table Grid'

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(*WHITE)
        shade_cell(cell, '1B3A6B')

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val) if val is not None else '')
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(*DGREY)
            shade_cell(cell, fill)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_footer(doc, project_name):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section_obj in doc.sections:
        footer = section_obj.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for text, bold, color in [
            ('Confidential  |  ', False, (0x80, 0x80, 0x80)),
            (project_name, True, NAVY),
            ('  |  Page ', False, (0x80, 0x80, 0x80)),
        ]:
            r = fp.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(8)
            r.bold = bold
            r.font.color.rgb = RGBColor(*color)

        # PAGE field
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.text = ' PAGE '
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run_xml = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), 'Calibri')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '16')
        rpr.append(rf)
        rpr.append(sz)
        run_xml.append(rpr)
        run_xml.append(fld)
        run_xml.append(instr)
        run_xml.append(fld2)
        fp._p.append(run_xml)


def main():
    if len(sys.argv) < 3:
        print('Usage: gen_docx.py <doc_data.json> <output.docx>', file=sys.stderr)
        sys.exit(1)
    check_import()

    from docx import Document
    from docx.shared import Cm, Pt

    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin  = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)

    meta    = data.get('meta', {})
    summary = data.get('summary', {})
    tables  = data.get('tables', [])
    project = meta.get('project_name', 'Power BI Model')

    try:
        dt = datetime.fromisoformat(meta.get('generated_at', '').replace('Z', '+00:00'))
        date_str = dt.strftime('%B %d, %Y')
    except Exception:
        date_str = meta.get('generated_at', '')[:10]

    # ── COVER PAGE ──────────────────────────────────────────────────────────
    for _ in range(8):
        doc.add_paragraph()

    add_para(doc, project.upper(), bold=True, size=36, color=NAVY, align='center', sa=4)
    add_para(doc, '━' * 48,        bold=False, size=11,  color=GOLD, align='center', sb=0, sa=4)
    add_para(doc, 'POWER BI SEMANTIC MODEL', bold=False, size=14, color=GOLD,  align='center', sa=2)
    add_para(doc, 'TECHNICAL DOCUMENTATION', bold=True,  size=14, color=NAVY, align='center', sa=0)

    for _ in range(10):
        doc.add_paragraph()

    add_para(doc, f'Prepared:  {date_str}',
             size=10, color=(0x60, 0x60, 0x60), align='center', sa=2)
    add_para(doc, f'Source:  {meta.get("pbip_dir", "Semantic Model")} ({meta.get("format","").upper()})',
             size=9, color=(0x80, 0x80, 0x80), align='center', sa=2)
    add_para(doc, 'C O N F I D E N T I A L',
             bold=True, size=9, color=(0x90, 0x90, 0x90), align='center')
    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────────────────────────────────
    add_heading(doc, '', 'Table of Contents', level=1)
    toc = [
        ('1.', 'Executive Summary'),
        ('2.', 'Data Model Overview'),
        ('3.', 'Table Details'),
        ('4.', 'Measures & KPIs'),
        ('5.', 'Business Logic'),
        ('6.', 'Data Sources'),
    ]
    if data.get('report_pages'):
        toc.append(('7.', 'Report Pages'))
        toc.append(('8.', 'Model Health'))
    else:
        toc.append(('7.', 'Model Health'))
    for num, title in toc:
        add_para(doc, f'    {num}   {title}', size=10, color=DGREY, sb=2, sa=2)
    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ────────────────────────────────────────────────
    add_heading(doc, '1.', 'Executive Summary')
    for chunk in (data.get('overview') or '').split('\n\n'):
        if chunk.strip():
            add_para(doc, chunk.strip(), size=10, color=DGREY, sb=4, sa=8)

    add_para(doc, 'At a Glance', bold=True, size=11, color=NAVY, sb=8, sa=4)
    add_table(doc,
        ['Metric', 'Value'],
        [
            ['Tables',        f'{summary.get("table_count",0)}  ({summary.get("fact_count",0)} fact · {summary.get("dimension_count",0)} dimension · {summary.get("other_count",0)} other)'],
            ['Measures',      f'{summary.get("measure_count",0)}  ({summary.get("documented_measures",0)} documented)'],
            ['Columns',       f'{summary.get("column_count",0)}  ({summary.get("calc_columns",0)} calculated · {summary.get("hidden_columns",0)} hidden)'],
            ['Relationships', f'{summary.get("relationship_count",0)}  ({summary.get("bidi_count",0)} bidirectional)'],
            ['Report Pages',  str(summary.get("report_pages", "N/A"))],
        ],
        col_widths=[5.5, 11]
    )
    doc.add_page_break()

    # ── 2. DATA MODEL OVERVIEW ───────────────────────────────────────────────
    add_heading(doc, '2.', 'Data Model Overview')

    add_heading(doc, '2.1', 'Entity Relationship Diagram', level=2)
    erd = data.get('erd', '')
    if erd:
        p = doc.add_paragraph()
        r = p.add_run(erd)
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)

    add_heading(doc, '2.2', 'Tables', level=2)
    if tables:
        add_table(doc,
            ['Table', 'Type', 'Columns', 'Measures', 'Description'],
            [[t.get('name',''), t.get('type','').title(),
              str(len(t.get('columns',[]))), str(len(t.get('measures',[]))),
              (t.get('description') or '')[:80]] for t in tables],
            col_widths=[3.5, 2.5, 2, 2.5, 6]
        )
    doc.add_page_break()

    # ── 3. TABLE DETAILS ─────────────────────────────────────────────────────
    add_heading(doc, '3.', 'Table Details')
    for ti, t in enumerate(tables):
        add_heading(doc, f'3.{ti+1}', f'{t.get("name","?")}  ({t.get("type","other").title()})', level=2)
        if t.get('description'):
            add_para(doc, t['description'], size=10, color=DGREY, sa=8)

        if t.get('columns'):
            add_para(doc, 'Columns', bold=True, size=10, color=NAVY, sb=6, sa=3)
            add_table(doc,
                ['Column', 'Data Type', 'Key', 'Hidden', 'Notes'],
                [[c.get('name',''), c.get('data_type',''),
                  'PK' if c.get('is_key') else '',
                  'Yes' if c.get('is_hidden') else '',
                  c.get('notes') or ''] for c in t['columns']],
                col_widths=[4, 3, 1.5, 2, 6]
            )

        if t.get('measures'):
            add_para(doc, 'Measures', bold=True, size=10, color=NAVY, sb=6, sa=3)
            add_table(doc,
                ['Measure', 'Display Folder', 'Format', 'Description'],
                [[m.get('name',''), m.get('folder') or '', m.get('format') or '',
                  (m.get('description') or '')[:65]] for m in t['measures']],
                col_widths=[4.5, 3, 2.5, 6.5]
            )
    doc.add_page_break()

    # ── 4. MEASURES & KPIs ───────────────────────────────────────────────────
    add_heading(doc, '4.', 'Measures & KPIs')
    from docx.shared import Pt, Cm, RGBColor
    folders = {}
    for t in tables:
        for m in t.get('measures', []):
            key = m.get('folder') or t.get('name', 'Other')
            folders.setdefault(key, []).append(m)

    for fi, (folder, measures) in enumerate(sorted(folders.items())):
        add_heading(doc, f'4.{fi+1}', folder, level=2)
        for m in measures:
            add_para(doc, m.get('name', ''), bold=True, size=10, color=NAVY, sb=10, sa=2)
            if m.get('description'):
                add_para(doc, m['description'], size=10, color=DGREY, sa=3)
            parts = []
            if m.get('format'):
                parts.append(f'Format: {m["format"]}')
            if parts:
                add_para(doc, '  ·  '.join(parts), size=9, color=(0x70, 0x70, 0x70), sa=3)
            if m.get('expression'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(6)
                r = p.add_run(m['expression'])
                r.font.name = 'Courier New'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(*CGREEN)
    doc.add_page_break()

    # ── 5. BUSINESS LOGIC ────────────────────────────────────────────────────
    add_heading(doc, '5.', 'Business Logic')
    for chunk in (data.get('business_logic') or '').split('\n\n'):
        if chunk.strip():
            add_para(doc, chunk.strip(), size=10, color=DGREY, sb=4, sa=8)
    doc.add_page_break()

    # ── 6. DATA SOURCES ──────────────────────────────────────────────────────
    add_heading(doc, '6.', 'Data Sources')
    sources = data.get('data_sources', [])
    if sources:
        add_table(doc,
            ['Table', 'Source Type', 'Connection'],
            [[s.get('table',''), s.get('type',''), s.get('connection','')] for s in sources],
            col_widths=[4, 3.5, 9]
        )
    else:
        add_para(doc, 'Source expressions not available in this project format.', size=10, color=(0x80,0x80,0x80))

    # ── 7. REPORT PAGES (optional) ───────────────────────────────────────────
    report_pages = data.get('report_pages', [])
    if report_pages:
        doc.add_page_break()
        add_heading(doc, '7.', 'Report Pages')
        add_table(doc,
            ['Page', 'Visuals', 'Visual Types'],
            [[p.get('name',''), str(p.get('visual_count',0)),
              ', '.join(p.get('visual_types',[]))] for p in report_pages],
            col_widths=[5, 2.5, 9]
        )

    # ── MODEL HEALTH (last section) ──────────────────────────────────────────
    doc.add_page_break()
    health_num = '8.' if report_pages else '7.'
    add_heading(doc, health_num, 'Model Health')

    health = data.get('health', {})
    if health.get('summary'):
        add_para(doc, health['summary'], size=10, color=DGREY, sa=8)

    health_rows = []
    if health.get('bidi_relationships'):
        health_rows.append(['Bidirectional relationships', ', '.join(health['bidi_relationships']), 'Review — may affect query performance'])
    if health.get('undocumented_count', 0) > 0:
        health_rows.append(['Undocumented measures', str(health['undocumented_count']), 'Run /pbi-comment-batch to add descriptions'])
    if health.get('missing_format_count', 0) > 0:
        health_rows.append(['Missing format strings', str(health['missing_format_count']), 'Add format strings for consistent display'])
    if health.get('unhidden_keys'):
        health_rows.append(['Unhidden key columns', ', '.join(health['unhidden_keys']), 'Consider hiding FK/surrogate key columns'])
    if health.get('isolated_tables'):
        health_rows.append(['Isolated tables', ', '.join(health['isolated_tables']), 'No relationships — review for orphaned tables'])

    if health_rows:
        add_table(doc, ['Issue', 'Detail', 'Recommendation'], health_rows, col_widths=[4, 4, 8.5])
    else:
        add_para(doc, 'No significant health concerns detected. Model follows best practices.', size=10, color=(0x1E,0x7A,0x1E), sb=8)

    # ── FOOTER ───────────────────────────────────────────────────────────────
    add_footer(doc, project)

    doc.save(sys.argv[2])
    print(f'DOCX_OK {sys.argv[2]}')


if __name__ == '__main__':
    main()
